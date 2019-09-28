# -*- coding: utf-8 -*-
"""
Created on Mon Mar 25 17:20:29 2019

@author: x103447
"""


#------convertDocxToText--------------------------------------------
import nltk, re
from docx import Document

def convertDocxToText(path):
	document = Document(path)
	return "\n".join([para.text for para in document.paragraphs])

abc = convertDocxToText(r"C:\Users\x103447\Desktop\Testing\SSIS\Python\Learn PY\NLP\Resume\neeraj goel.docx")
print(abc)


#############################################################
#Get Name
#############################################################
lines = [el.strip() for el in abc.split("\n") if len(el) > 0]
lines = [nltk.word_tokenize(el) for el in lines]
lines = [nltk.pos_tag(el) for el in lines]
#print(lines)
grammar = r'NAME: {<NN.*>*}'
#grammar = r'NAME: {<NN.*><NN.*><NN.*>*}'
chunkParser = nltk.RegexpParser(grammar)
# Reads Indian Names from the file, reduce all to lower case for easy comparision [Name lists]
indianNames = open("allNames.txt", "r").read().lower()
# Lookup in a set is much faster
indianNames = set(indianNames.split())
nameHits = []
otherNameHits = []
for tagged_tokens in lines:
    chunked_tokens = chunkParser.parse(tagged_tokens)
    for subtree in chunked_tokens.subtrees():
        if subtree.label() == 'NAME':
            for ind, leaf in enumerate(subtree.leaves()):
                if leaf[0].lower() in indianNames and 'NN' in leaf[1]:
#                    print(leaf)
#                    for el in subtree.leaves()[ind:ind+3]:
#                        print(el[0])
#                        hit = " ".join(el[0])
                    hit = " ".join([el[0] for el in subtree.leaves()[ind:ind+3]])
                    nameHits.append(hit)

# Going for the first name hit
if len(nameHits) > 0:
#    print(nameHits[0])
    nameHits = [re.sub(r'[^a-zA-Z \-]', '', el).strip() for el in nameHits] 
    name = " ".join([el[0].upper()+el[1:].lower() for el in nameHits[0].split() if len(el)>0])
    otherNameHits = nameHits[1:]

print(nameHits)    
print(otherNameHits)    
print(name) #Final Name result

#############################################################
#Get email id
#############################################################
#print(abc)
inputString = abc

email = None
try:
    pattern = re.compile(r'\S*@\S*')
    matches = pattern.findall(inputString) # Gets all email addresses as a list
    email = matches
#    print("found")
except Exception as e:
#    print("Not found")
    print (e)
print (email)

#############################################################
#get phone number  
#############################################################
inputString = abc          
number = None
try:
    pattern = re.compile(r'([+(]?\d+[)\-]?[ \t\r\f\v]*[(]?\d{2,}[()\-]?[ \t\r\f\v]*\d{2,}[()\-]?[ \t\r\f\v]*\d*[ \t\r\f\v]*\d*[ \t\r\f\v]*)')
        # Understanding the above regex
        # +91 or (91) -> [+(]? \d+ -?
        # Metacharacters have to be escaped with \ outside of character classes; inside only hyphen has to be escaped
        # hyphen has to be escaped inside the character class if you're not incidication a range
        # General number formats are 123 456 7890 or 12345 67890 or 1234567890 or 123-456-7890, hence 3 or more digits
        # Amendment to above - some also have (0000) 00 00 00 kind of format
        # \s* is any whitespace character - careful, use [ \t\r\f\v]* instead since newlines are trouble
    match = pattern.findall(inputString)
    # match = [re.sub(r'\s', '', el) for el in match]
        # Get rid of random whitespaces - helps with getting rid of 6 digits or fewer (e.g. pin codes) strings
    # substitute the characters we don't want just for the purpose of checking
    match = [re.sub(r'[,.]', '', el) for el in match if len(re.sub(r'[()\-.,\s+]', '', el))>6]
        # Taking care of years, eg. 2001-2004 etc.
    match = [re.sub(r'\D$', '', el).strip() for el in match]
        # $ matches end of string. This takes care of random trailing non-digit characters. \D is non-digit characters
    match = [el for el in match if len(re.sub(r'\D','',el)) <= 15]
        # Remove number strings that are greater than 15 digits
    try:
        for el in list(match):
            # Create a copy of the list since you're iterating over it
            if len(el.split('-')) > 3: continue # Year format YYYY-MM-DD
            for x in el.split("-"):
                try:
                    # Error catching is necessary because of possibility of stray non-number characters
                    # if int(re.sub(r'\D', '', x.strip())) in range(1900, 2100):
                    if x.strip()[-4:].isdigit():
                        if int(x.strip()[-4:]) in range(1900, 2100):
                            # Don't combine the two if statements to avoid a type conversion error
                            match.remove(el)
                except:
                    pass
    except:
        pass
    number = match
except:
    pass
print(number)


#############################################################
# get Experience details
#############################################################
sen = ""
experience = []
lines = [el.strip() for el in abc.split("\n") if len(el) > 0]
lines = [nltk.word_tokenize(el) for el in lines]
lines = [nltk.pos_tag(el) for el in lines]
try:
    for sentence in lines:#find the index of the sentence where the degree is find and then analyse that sentence
            sen=" ".join([words[0].lower() for words in sentence]) #string of words in sentence
            if re.search('experience',sen):
                sen_tokenised= nltk.word_tokenize(sen)
                tagged = nltk.pos_tag(sen_tokenised)
                entities = nltk.chunk.ne_chunk(tagged)
                for subtree in entities.subtrees():
                    for leaf in subtree.leaves():
                        if leaf[1]=='CD':
#                            experience=leaf[0]
                            experience.append(leaf[0])
except Exception as e:
#    print (traceback.format_exc())
    print (e)

print(experience)
print(max(experience)) #if list has multiple experience


#Extracting Email_id and Phone numbers using pattern
pattern = re.compile(r'([+()]?\d+[)\-]?[ \t\r\f\v]*[(]?\d{2,}[()\-]?[ \t\r\f\v]*\d{2,}[()\-]?[ \t\r\f\v]*\d*[ \t\r\f\v]*\d*[ \t\r\f\v]*)')


#############################################################
#------convertPDFToText--------------------------------------
#############################################################
import io
from pdfminer.pdfinterp import PDFResourceManager, PDFPageInterpreter
from pdfminer.converter import TextConverter
from pdfminer.layout import LAParams
from pdfminer.pdfpage import PDFPage

def convertPDFToText(path):
    #to store the shared resources like font or images
    rsrcmgr = PDFResourceManager()
    #to read string either of 8 bit or unicode
    retstr = io.StringIO()
    #UTF stands for Unicode Transformation Format. The '8' means it uses 8-bit blocks to represent a character
    codec = 'utf-8'
    # to perform layout analysis like LTTextBox LTTextLine
    laparams = LAParams()
    #to transalte it to whatever you need
    device = TextConverter(rsrcmgr, retstr, codec=codec, laparams=laparams)
    #to Open the file in read binary mode
    fp = io.FileIO(path, 'rb')
    #to process the page content 
    interpreter = PDFPageInterpreter(rsrcmgr, device)
    password = ""
    maxpages = 0
    caching = True
    pagenos=set() 
    
    for page in PDFPage.get_pages(fp, pagenos, maxpages=maxpages, password=password,caching=caching, check_extractable=True):
        interpreter.process_page(page)
    fp.close()
    device.close()
    string = retstr.getvalue()
    retstr.close()
    return string


abc = convertPDFToText(r"C:\Users\x103447\Desktop\Testing\SSIS\Python\Learn PY\NLP\Resume\atul sharma.pdf")
print(abc)


#------convertRtfToText--------------------------------------------

from pyth.plugins.rtf15.reader import Rtf15Reader
from pyth.plugins.plaintext.writer import PlaintextWriter
import pyth 
#print(pyth.__file__)
#pyth.plugins.rtf15.reader.Rtf15Reader
#pyth.plugins.rtf15.reader.Rtf15Reader.read()

def convertRtfToText(path):
	doc = pyth.plugins.rtf15.reader.Rtf15Reader.read(open(path))
	return pyth.plugins.plaintext.writer.PlaintextWriter.write(doc).getvalue()

abc = convertRtfToText(r"C:\Users\x103447\Desktop\Testing\SSIS\Python\Learn PY\NLP\Resume\atul sharma.pdf")


#abc = (r"C:\Users\x103447\Desktop\Testing\SSIS\Python\Learn PY\NLP\Resume\neeraj goel.docx")
#document = Document(abc)
#
#lines = [el.strip() for el in abc.split("\n") if len(el) > 0]
#lines = [nltk.word_tokenize(el) for el in lines]
#lines = [nltk.pos_tag(el) for el in lines]
#
#sentences = nltk.sent_tokenize(abc)
#sentences =[nltk.word_tokenize(sent) for sent in sentences]
#tokens = sentences
#sentences = [nltk.pos_tag(sent) for sent in sentences]
#dummy = []
#for el in tokens:
#    dummy += el
#tokens = dummy    


#------convertDocxToText--------------------------------------------
#from docx import Document
#from docx.shared import Inches
#
#document = Document()
#
#document.add_heading('Document Title', 0)
#
#p = document.add_paragraph('A plain paragraph having some ')
#p.add_run('bold').bold = True
#p.add_run(' and some ')
#p.add_run('italic.').italic = True
#
#document.add_heading('Heading, level 1', level=1)
#document.add_paragraph('Intense quote', style='Intense Quote')
#
#document.add_paragraph(
#    'first item in unordered list', style='List Bullet')
#document.add_paragraph(
#    'first item in ordered list', style='List Number')
#
#records = (
#    (3, '101', 'Spam'),
#    (7, '422', 'Eggs'),
#    (4, '631', 'Spam, spam, eggs, and spam')
#)
#
#table = document.add_table(rows=1, cols=3)
#hdr_cells = table.rows[0].cells
#hdr_cells[0].text = 'Qty'
#hdr_cells[1].text = 'Id'
#hdr_cells[2].text = 'Desc'
#for qty, id, desc in records:
#    row_cells = table.add_row().cells
#    row_cells[0].text = str(qty)
#    row_cells[1].text = id
#    row_cells[2].text = desc
#
#document.add_page_break()
#
#document.save('demo.docx')
